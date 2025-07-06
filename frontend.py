import streamlit as st
import requests
from docx import Document
from io import BytesIO

st.title("📝 Auto-Quiz Generator (Ollama)")

# File upload
uploaded_file = st.file_uploader("Upload Textbook PDF", type="pdf")

# Number of questions input
num_questions = st.number_input("Number of Questions", min_value=1, max_value=20, value=3)

# Question type select
question_type = st.selectbox("Question Type", ["mcq", "truefalse"])

if uploaded_file:
    if st.button("Generate Quiz"):
        with st.spinner("Generating quiz..."):
            try:
                response = requests.post(
                    "http://localhost:5001/generate_quiz",
                    files={"file": uploaded_file},
                    data={
                        "num_questions": num_questions,
                        "question_type": question_type
                    }
                )
                response.raise_for_status()
                quiz = response.json().get("quiz", "No quiz returned.")

                st.write("### Generated Quiz")
                st.text(quiz)

                # ✅ Generate DOCX
                doc = Document()
                doc.add_heading("Generated Quiz", 0)
                for line in quiz.split("\n"):
                    doc.add_paragraph(line)

                # ✅ Save to in-memory buffer
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # ✅ Download button
                st.download_button(
                    label="📄 Download Quiz as DOCX",
                    data=buffer,
                    file_name="generated_quiz.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"Failed to generate quiz: {e}")
